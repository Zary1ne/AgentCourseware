# -*- coding: utf-8 -*-
"""
AI 教学助手 - 基于 LangChain 架构
课程参考：
- LangChain入门-初始化模型
- LangChain入门-模型-调用模型
- LangChain入门-消息 (Message)
- LangChain入门-提示词工程
- LangChain入门-短期记忆
"""

import json, os, re, sys, time, io

# 加载 .env 配置文件
try:
    from dotenv import load_dotenv
    env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')
    if os.path.exists(env_path):
        load_dotenv(env_path)
        print(f"[OK] 已加载 .env 配置文件: {env_path}")
    else:
        print(f"[WARN] .env 文件不存在: {env_path}")
except ImportError:
    print("[WARN] python-dotenv 未安装")

if sys.platform == "win32":
    try: sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except: pass

from flask import Flask, request, Response, send_from_directory, send_file

# ==================== 配置项 ====================
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com")
DEEPSEEK_MODEL = os.getenv("DEEPSEEK_MODEL", "deepseek-chat")
FLASK_HOST = os.getenv("FLASK_HOST", "0.0.0.0")
FLASK_PORT = int(os.getenv("FLASK_PORT", "8000"))

if not DEEPSEEK_API_KEY:
    print("[ERROR] 未找到 DEEPSEEK_API_KEY，请在 .env 文件中配置")
    sys.exit(1)

app = Flask(__name__, static_folder=".", static_url_path="")

# ==================== 1. 初始化模型 ====================
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain_community.chat_message_histories import ChatMessageHistory

llm = ChatOpenAI(
    model=DEEPSEEK_MODEL,
    openai_api_key=DEEPSEEK_API_KEY,
    openai_api_base=DEEPSEEK_BASE_URL + "/v1",
    streaming=True,
    temperature=0.7,
    max_tokens=2048,
)
print("[OK] LangChain 模型初始化完成")

# ==================== 2. 短期记忆 ====================
SESSION_HISTORY = {}

def get_session_history(session_id):
    if session_id not in SESSION_HISTORY:
        SESSION_HISTORY[session_id] = ChatMessageHistory()
    return SESSION_HISTORY[session_id]

def clear_session_history(session_id):
    if session_id in SESSION_HISTORY:
        SESSION_HISTORY[session_id] = ChatMessageHistory()

# ==================== 3. 提示词工程 ====================
SYSTEM_PROMPT = """你是全科学习助手，可以帮助学生学习各科知识。

支持的科目包括但不限于：
- 数学：代数、几何、微积分、概率统计等
- 语文：古诗文、现代文阅读、写作、修辞手法等
- 英语：语法、词汇、阅读理解、写作、翻译等
- 物理：力学、电磁学、光学、热学等
- 化学：有机化学、无机化学、反应原理等
- 生物：细胞、遗传、生态等
- 历史：中国历史、世界历史等
- 地理：自然地理、人文地理等
- 政治：经济、政治、哲学、文化等
- 计算机：编程、算法、数据结构等
- 其他任何学科知识

你的特点：
- 根据科目特点用最合适的方式讲解
- 数学用公式和例题，语文用赏析和分析，英语用语法和例句
- 用最简单的话解释概念，避免术语堆砌
- 回复始终用中文（专业术语除外）

对话原则：
1. 识别用户问的是哪个科目
2. 闲聊时自然回应，不强行输出教案
3. 问具体知识点时，用「概念解释 + 例子/例题 + 练习建议」三段式回答
4. 记住之前的对话上下文，保持对话连贯性

格式要求：
- 禁止使用 * 字符做格式标记
- 用数字 1. 2. 3. 或中文标点做列表
- 代码或公式用正确的格式展示
- 直接给答案，不要以"好的，我来帮你..."开头
- 始终用中文回复"""

TEACHING_PLAN_PROMPT = '''你是教学设计专家。请根据用户的学科和主题，生成一份完整的教学方案。

必须包含以下部分，用【】标注标题：

【教学目标】
- 学生学完能做什么（具体的、可验证的目标）
- 涉及的核心概念有哪些

【教学重难点】
- 重点：必须掌握的核心内容
- 难点：容易卡住的地方及突破方法

【教学过程】
- 导入环节（如何引入主题）
- 新课讲解（主要知识点讲解）
- 例题分析（典型例题或案例）
- 课堂练习（随堂练习）

【课后作业】
- 基础题（2-3道，巩固基础）
- 提高题（1-2道，拓展思维）
- 思考题（1道，联系实际）

要求：
1. 内容紧扣用户的学科和主题
2. 根据学科特点选择合适的教学方法
3. 数学/物理等理科：公式、例题、计算步骤
4. 语文/英语等文科：赏析、分析、写作指导
5. 语言简单直白，适合学生
6. 禁止使用 * 字符做任何格式标记
7. 始终用中文回复'''

ANALYSIS_PROMPT = """你是资深教学解析专家。当用户请求教学解析时，请按以下结构回答：

1. 【教学定位】这个知识点在课程体系中的位置和作用
2. 【重点分析】必须掌握的核心内容和关键概念
3. 【难点突破】学生容易卡住的地方，以及突破方法
4. 【教学建议】教学顺序、教学方法建议
5. 【常见误区】学生常犯的错误和纠正方法
6. 【拓展延伸】相关知识点关联和拓展方向

要求：
- 结合具体学科特点
- 语言通俗易懂
- 禁止使用 * 字符做格式标记"""

SUGGESTIONS_PROMPT = """你是个性化学习顾问。当用户请求学习建议时，请按以下结构回答：

1. 【学习诊断】分析学生的学习现状和可能的薄弱点
2. 【学习目标】制定明确、可衡量的学习目标
3. 【学习计划】分阶段的学习计划（周计划/月计划）
4. 【学习方法】适合该学科的高效学习方法和技巧
5. 【练习推荐】针对性的练习题和学习资源
6. 【进度追踪】如何检验学习效果、调整计划

要求：
- 方案具体可行
- 方法科学有效
- 禁止使用 * 字符做格式标记"""

EXPLAIN_PROMPT = """你是知识讲解专家。当用户请求知识讲解时，请按以下结构回答：

1. 【概念定义】用最简洁的语言定义这个知识点
2. 【通俗解释】用生活中的例子或比喻来解释
3. 【原理分析】深入讲解背后的原理和逻辑
4. 【示例演示】给出典型的例题或应用场景
5. 【注意事项】使用时的注意点和常见陷阱
6. 【举一反三】相关知识点的关联和延伸

要求：
- 由浅入深，循序渐进
- 多用直观的例子
- 禁止使用 * 字符做格式标记"""

ENGLISH_PROMPT = """你是高中英语专家，专注于高一英语教学。请根据高一英语教学大纲，提供以下内容：

1. 【词汇学习】重点词汇、短语辨析、用法举例
2. 【语法讲解】高一核心语法点（时态、从句、非谓语动词等）
3. 【阅读理解】阅读技巧、常见题型分析
4. 【写作指导】写作框架、常用句型、范文赏析
5. 【听力训练】听力技巧、常见场景词汇
6. 【口语表达】日常交际用语、情景对话

要求：
- 内容贴合高一水平
- 例句实用地道
- 中英双语解释
- 禁止使用 * 字符做格式标记"""

# ==================== 4. 构建对话链 ====================
PROMPT_MAP = {
    "general": SYSTEM_PROMPT,
    "teaching": TEACHING_PLAN_PROMPT,
    "analysis": ANALYSIS_PROMPT,
    "suggestions": SUGGESTIONS_PROMPT,
    "explain": EXPLAIN_PROMPT,
    "english": ENGLISH_PROMPT,
}

def build_conversation_chain(system_prompt):
    prompt = ChatPromptTemplate.from_messages([
        ("system", system_prompt),
        ("human", "{input}"),
    ])
    chain = prompt | llm
    return RunnableWithMessageHistory(
        chain,
        get_session_history=get_session_history,
        input_messages_key="input",
        history_messages_key="history",
    )

CHAIN_MAP = {k: build_conversation_chain(v) for k, v in PROMPT_MAP.items()}

# ==================== 5. 导出功能 ====================
try:
    from pptx_builder import create_pptx, parse_content_to_slides
    HAS_PPTX = True
except:
    HAS_PPTX = False

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except:
    HAS_DOCX = False

# ==================== 6. 工具函数 ====================
def clean_response(text):
    return text.replace("*", "")

def detect_chain_type(text, prompt_type=""):
    if prompt_type and prompt_type in CHAIN_MAP:
        return prompt_type
    teaching_keywords = ["教学目标", "教学设计", "教案", "备课", "教学方案", "重难点"]
    if any(k in text for k in teaching_keywords):
        return "teaching"
    return "general"

def stream_fallback(text):
    for ch in clean_response(text):
        yield f"data: {json.dumps({'type': 'content', 'data': ch}, ensure_ascii=False)}\n\n"
        time.sleep(0.015)
    yield f"data: {json.dumps({'type': 'done'}, ensure_ascii=False)}\n\n"

# ==================== 7. SSE 流式输出 ====================
def stream_chat(user_text, session_id, chain_type):
    try:
        chain = CHAIN_MAP.get(chain_type, CHAIN_MAP["general"])
        stream_iter = chain.stream(
            {"input": user_text},
            config={"configurable": {"session_id": session_id}}
        )
        full_response = ""
        for chunk in stream_iter:
            if hasattr(chunk, 'content') and chunk.content:
                text = clean_response(chunk.content)
                full_response += text
                for ch in text:
                    yield f"data: {json.dumps({'type': 'content', 'data': ch}, ensure_ascii=False)}\n\n"
                    time.sleep(0.012)
        print(f"[OK] 回复完成，长度: {len(full_response)} 字符")
        yield f"data: {json.dumps({'type': 'done'}, ensure_ascii=False)}\n\n"
    except Exception as e:
        print(f"[ERROR] API 调用失败: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        error_msg = f"抱歉，AI 服务暂时不可用，请稍后重试。\n\n[错误: {type(e).__name__}: {e}]"
        for chunk in stream_fallback(error_msg):
            yield chunk

# ==================== 8. DOCX 构建 ====================
def build_docx(title, content):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    tp = doc.add_heading(title, level=0)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in tp.runs: r.font.size = Pt(18)
    for line in content.strip().split("\n"):
        line = line.strip()
        if not line: continue
        if line.startswith("【") and "】" in line:
            h = doc.add_heading(line, level=2)
            for r in h.runs: r.font.size = Pt(14)
        elif re.match(r"^\d+[\.\、\)]", line):
            doc.add_paragraph(line, style="List Number")
        elif re.match(r"^[•\-—]", line):
            doc.add_paragraph(line[1:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ==================== 9. 路由 ====================
@app.route("/")
def index():
    return send_from_directory(".", "index.html")

@app.route("/api/v1/health")
def health():
    return {
        "status": "ok",
        "api_configured": bool(DEEPSEEK_API_KEY),
        "langchain": True,
        "export_docx": HAS_DOCX,
        "export_pptx": HAS_PPTX
    }

@app.route("/api/v1/new-session", methods=["POST"])
def new_session():
    data = request.get_json(force=True)
    session_id = data.get("session_id", "default_session")
    clear_session_history(session_id)
    return {"status": "ok", "message": "会话已重置"}

@app.route("/api/v1/chat", methods=["POST"])
def chat():
    data = request.get_json(force=True)
    user_text = data.get("message", "").strip()
    session_id = data.get("session_id", "default_session")
    prompt_type = data.get("prompt_type", "")
    
    if not user_text:
        return Response(
            stream_fallback("请输入你的问题。"),
            mimetype="text/event-stream",
            headers={"Cache-Control": "no-cache", "Connection": "keep-alive"}
        )

    chain_type = detect_chain_type(user_text, prompt_type)
    
    return Response(
        stream_chat(user_text, session_id, chain_type),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )

@app.route("/api/v1/export/docx", methods=["POST"])
def export_docx():
    if not HAS_DOCX: return {"error": "python-docx not available"}, 500
    data = request.get_json(force=True)
    content = data.get("content", "")
    title = data.get("title", "教学文档")
    if not content: return {"error": "No content"}, 400
    try:
        buf = build_docx(title, content)
        safe = re.sub(r'[\\/*?:"<>|]', '_', title)
        return send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"{safe}.docx"
        )
    except Exception as e:
        return {"error": str(e)}, 500

@app.route("/api/v1/export/pptx", methods=["POST"])
def export_pptx():
    if not HAS_PPTX: return {"error": "PPTX builder not available"}, 500
    data = request.get_json(force=True)
    content = data.get("content", "")
    title = data.get("title", "教学课件")
    if not content: return {"error": "No content"}, 400
    try:
        slides = parse_content_to_slides(content)
        pptx_bytes = create_pptx(title, slides)
        safe = re.sub(r'[\\/*?:"<>|]', '_', title)
        return send_file(
            io.BytesIO(pptx_bytes),
            mimetype="application/vnd.openxmlformats-officedocument.presentationml.presentation",
            as_attachment=True,
            download_name=f"{safe}.pptx"
        )
    except Exception as e:
        return {"error": str(e)}, 500

if __name__ == "__main__":
    print(f"[OK] DeepSeek API: {'已配置' if DEEPSEEK_API_KEY else '未配置'}")
    print(f"[OK] LangChain: 可用")
    print(f"[OK] Word导出: {'可用' if HAS_DOCX else '不可用'}")
    print(f"[OK] PPT导出: {'可用' if HAS_PPTX else '不可用'}")
    print(f"[INFO] http://localhost:{FLASK_PORT}")
    app.run(host=FLASK_HOST, port=FLASK_PORT, debug=False, threaded=True)
