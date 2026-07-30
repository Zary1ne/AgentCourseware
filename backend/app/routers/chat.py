import asyncio
import json
import uuid
import re
import io
import time
from urllib.parse import quote
from fastapi import APIRouter, HTTPException, Request
from fastapi.responses import StreamingResponse, Response
from sse_starlette.sse import EventSourceResponse
from app.models.schemas import ChatRequest, ChatMessage
from app.services.llm_service import chat_stream, chat_sync, extract_intent, clean_response
from app.services.rag_service import search_knowledge
from app.routers.admin import record_conversation

router = APIRouter(prefix="/api/chat", tags=["chat"])

# 存储会话（简单内存存储，生产环境应使用Redis）
sessions = {}

# ===== 关键词意图识别（来自 AITEACH(对话) backend/main.py 原型） =====
SUBJECT_KEYWORDS = {
    "物理": ["物理", "力学", "电学", "牛顿", "运动", "能量", "光的", "磁"],
    "数学": ["数学", "函数", "几何", "代数", "方程", "概率", "微积分", "三角"],
    "化学": ["化学", "元素", "反应", "分子", "原子", "酸碱", "氧化", "有机"],
    "生物": ["生物", "细胞", "基因", "光合", "生态", "遗传", "进化", "DNA"],
    "语文": ["语文", "古诗", "文言文", "阅读", "作文", "修辞", "诗歌", "散文"],
    "英语": ["英语", "语法", "词汇", "阅读理解", "完形填空", "听力"],
    "历史": ["历史", "朝代", "战争", "革命", "古代", "近代", "文明"],
    "地理": ["地理", "气候", "地形", "河流", "板块", "人口", "城市"],
}

GRADE_KEYWORDS = {
    "高一": ["高一", "高中一年级"],
    "高二": ["高二", "高中二年级"],
    "高三": ["高三", "高中三年级"],
    "初一": ["初一", "七年级"],
    "初二": ["初二", "八年级"],
    "初三": ["初三", "九年级"],
}

TYPE_KEYWORDS = {
    "新课": ["新课", "新授课", "讲授"],
    "复习课": ["复习", "回顾", "总结课"],
    "实验课": ["实验", "动手", "操作"],
    "习题课": ["习题", "练习", "题", "做题"],
}


def recognize_intent(text: str) -> dict:
    """从用户输入中识别教学意图参数（关键词匹配）"""
    intents = {}
    for subject, keywords in SUBJECT_KEYWORDS.items():
        if any(kw in text for kw in keywords):
            intents["subject"] = subject
            break
    for grade, keywords in GRADE_KEYWORDS.items():
        if any(kw in text for kw in keywords):
            intents["grade"] = grade
            break
    for lesson_type, keywords in TYPE_KEYWORDS.items():
        if any(kw in text for kw in keywords):
            intents["type"] = lesson_type
            break
    duration_match = re.search(r"(\d+)\s*分钟", text)
    if duration_match:
        intents["duration"] = duration_match.group(1) + "分钟"
    if "严谨" in text or "学术" in text:
        intents["style"] = "学术严谨"
    elif "活泼" in text or "互动" in text or "有趣" in text:
        intents["style"] = "活泼互动"
    elif "简洁" in text or "商务" in text or "极简" in text:
        intents["style"] = "极简商务"
    return intents


def generate_ai_response(text: str, params: dict) -> str:
    """API 不可用时的硬编码兜底回复（来自 AITEACH(对话) backend/main.py 原型）"""
    subject = params.get("subject", "该学科")
    grade = params.get("grade", "相应年级")
    duration = params.get("duration", "45分钟")
    lesson_type = params.get("type", "新课")

    if "牛顿" in text or "物理" in text:
        return (
            f"好的！我来帮你设计一节{grade}{subject}——牛顿第一定律的{lesson_type}教学方案，课时{duration}。\n\n"
            "以下是结构化的教学大纲：\n\n"
            "【教学目标】\n"
            "· 知识与技能：理解牛顿第一定律的内容，知道惯性的概念\n"
            "· 过程与方法：通过理想实验法，培养科学推理能力\n"
            "· 情感态度价值观：体会科学探索精神，认识物理与生活的联系\n\n"
            "【教学重难点】\n"
            "· 重点：牛顿第一定律的内容及理解\n"
            "· 难点：惯性的概念及其与质量的关系\n\n"
            "【教学过程】\n"
            "1. 引入（5分钟）：通过生活实例（公交车急刹车）引出问题\n"
            "2. 讲解（15分钟）：伽利略理想斜面实验 → 笛卡尔补充 → 牛顿总结\n"
            "3. 实验（10分钟）：惯性演示实验（小车+木块突然运动）\n"
            "4. 练习（10分钟）：典型例题分析，判断物体运动状态\n"
            "5. 总结（5分钟）：知识梳理，布置课后思考题\n\n"
            "【所需教具】斜面、小车、木块、多媒体课件\n\n"
            "需要我调整哪个部分，或者直接生成课件吗？"
        )

    if "实验" in text:
        return (
            "收到！我注意到你提到了实验环节，这很重要。以下是包含实验设计的建议：\n\n"
            "【实验设计建议】\n"
            f"· 实验类型：演示实验 + 分组探究\n"
            f"· 建议时长：{duration}中安排15-20分钟实验时间\n"
            "· 安全提示：提前检查器材，强调操作规范\n\n"
            "【实验流程】\n"
            "1. 提出问题（2分钟）\n"
            "2. 学生猜想（3分钟）\n"
            "3. 动手实验（10分钟）\n"
            "4. 数据记录与分析（3分钟）\n"
            "5. 结论汇报（2分钟）\n\n"
            "你可以告诉我具体学科和知识点，我来生成完整的教学方案。"
        )

    if "复习" in text:
        return (
            "好的！复习课的设计需要注重知识梳理和查漏补缺。\n\n"
            "【复习课教学框架】\n"
            "1. 知识回顾（10分钟）：思维导图梳理本章核心概念\n"
            "2. 易错点辨析（10分钟）：典型错误案例分析\n"
            "3. 综合练习（15分钟）：分层练习题，从基础到提高\n"
            "4. 方法总结（5分钟）：解题技巧与思路归纳\n"
            "5. 课后任务（5分钟）：针对性练习布置\n\n"
            "请告诉我具体的学科和知识点，我可以生成更详细的复习方案。"
        )

    return (
        f"好的！我已经识别到你的教学需求：\n\n"
        f"· 学科：{subject}\n"
        f"· 年级：{grade}\n"
        f"· 课型：{lesson_type}\n"
        f"· 课时：{duration}\n\n"
        "接下来我会基于以上信息为你生成教学方案。你可以进一步补充：\n"
        "1. 具体的知识点或章节\n"
        "2. 是否需要实验环节\n"
        "3. 学生的学习基础情况\n"
        "4. 特别的教学偏好\n\n"
        "请补充以上信息，我将为你生成详细的教学大纲。"
    )


@router.post("/message")
async def send_message(request: ChatRequest):
    """发送聊天消息（流式响应）"""
    # 将消息历史转为字典列表
    history = [{"role": m.role, "content": m.content} for m in request.messages]

    # 最后一条用户消息用于检索知识库
    user_messages = [m for m in history if m["role"] == "user"]
    rag_context = ""
    if user_messages:
        last_user_msg = user_messages[-1]["content"]
        try:
            search_results = search_knowledge(last_user_msg, top_k=3)
            if search_results:
                rag_context = "\n\n【知识库参考资料】\n" + "\n---\n".join(
                    [f"来源: {r['source']}\n{r['content'][:500]}" for r in search_results]
                )
        except Exception:
            pass

    # 如果有RAG上下文，添加到系统提示中
    if rag_context:
        history = [
            {"role": "system", "content": f"以下是知识库中与当前教学主题相关的参考资料，请在回答时参考这些内容：{rag_context}"}
        ] + history

    # 记录使用
    record_conversation()

    if request.stream:
        async def event_generator():
            user_msgs = [m for m in history if m["role"] == "user"]
            last_user_text = user_msgs[-1]["content"] if user_msgs else ""
            keyword_intents = recognize_intent(last_user_text)
            if keyword_intents:
                yield {"event": "intents", "data": json.dumps(keyword_intents, ensure_ascii=False)}

            full_response = ""
            buffer = ""
            flush_interval = 0.05
            last_flush = time.monotonic()

            try:
                async for chunk in chat_stream(history, request.prompt_type):
                    full_response += chunk
                    buffer += chunk
                    now = time.monotonic()
                    if now - last_flush >= flush_interval or len(buffer) >= 50:
                        yield {"event": "message", "data": buffer}
                        buffer = ""
                        last_flush = now
                if buffer:
                    yield {"event": "message", "data": buffer}
            except Exception:
                fallback = generate_ai_response(last_user_text, keyword_intents)
                full_response = clean_response(fallback)
                yield {"event": "message", "data": full_response}

            if "[INTENT_READY]" in full_response:
                try:
                    intent = extract_intent(history + [{"role": "assistant", "content": full_response}])
                    yield {"event": "intent", "data": json.dumps(intent, ensure_ascii=False)}
                except Exception:
                    pass

            yield {"event": "done", "data": "[DONE]"}

        return EventSourceResponse(event_generator())

    else:
        # 非流式响应
        try:
            response_text = chat_sync(history, request.prompt_type)

            # 尝试提取意图
            intent = None
            if "[INTENT_READY]" in response_text:
                try:
                    intent = extract_intent(history + [{"role": "assistant", "content": response_text}])
                except Exception:
                    pass

            return {
                "message": response_text,
                "intent": intent,
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"LLM调用失败: {str(e)}")


@router.post("/extract-intent")
async def extract_teaching_intent(request: ChatRequest):
    """从对话历史中提取教学意图"""
    history = [{"role": m.role, "content": m.content} for m in request.messages]
    try:
        intent = extract_intent(history)
        return {"success": True, "intent": intent}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"意图提取失败: {str(e)}")


@router.get("/session/{session_id}")
async def get_session(session_id: str):
    """获取会话信息"""
    if session_id not in sessions:
        return {"session_id": session_id, "messages": []}
    return {"session_id": session_id, "messages": sessions[session_id]}


@router.post("/session")
async def create_session():
    """创建新会话"""
    session_id = uuid.uuid4().hex[:12]
    sessions[session_id] = []
    return {"session_id": session_id}


# ===== 导出端点（来自 AITEACH(对话) app.py 原型） =====
def build_docx(title: str, content: str) -> io.BytesIO:
    """从对话内容构建简易 Word 文档"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        HAS_DOCX = True
    except ImportError:
        HAS_DOCX = False

    if not HAS_DOCX:
        raise HTTPException(status_code=500, detail="python-docx not available")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    tp = doc.add_heading(title, level=0)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in tp.runs:
        r.font.size = Pt(18)
    for line in content.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("【") and "】" in line:
            h = doc.add_heading(line, level=2)
            for r in h.runs:
                r.font.size = Pt(14)
        elif re.match(r"^\d+[\.\、\)]", line):
            doc.add_paragraph(line, style="List Number")
        elif re.match(r"^[·•\-—]", line):
            doc.add_paragraph(line[1:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@router.post("/export/docx")
async def export_docx(request: Request):
    """从对话内容导出 Word 文档（来自 AITEACH(对话) 原型）"""
    data = await request.json()
    content = data.get("content", "")
    title = data.get("title", "教学文档")
    if not content:
        raise HTTPException(status_code=400, detail="No content")
    try:
        buf = build_docx(title, content)
        safe = re.sub(r'[\\/*?:"<>|]', '_', title)
        quoted = quote(safe, safe='')
        return Response(
            buf.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quoted}.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/export/pptx")
async def export_pptx(request: Request):
    """从对话内容导出 PPT 课件（来自 AITEACH(对话) 原型）"""
    try:
        from app.services.pptx_builder import create_pptx, parse_content_to_slides
        HAS_PPTX = True
    except ImportError:
        HAS_PPTX = False

    if not HAS_PPTX:
        raise HTTPException(status_code=500, detail="PPTX builder not available")

    data = await request.json()
    content = data.get("content", "")
    title = data.get("title", "教学课件")
    if not content:
        raise HTTPException(status_code=400, detail="No content")
    try:
        slides = parse_content_to_slides(content)
        pptx_bytes = create_pptx(title, slides)
        safe = re.sub(r'[\\/*?:"<>|]', '_', title)
        quoted = quote(safe, safe='')
        return Response(
            pptx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quoted}.pptx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
