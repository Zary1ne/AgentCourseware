import os
import uuid
import json
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from app.config import OUTPUT_DIR
from app.services.llm_service import generate_content


def create_teaching_plan(intent: dict, context: str = "") -> str:
    """生成 Word 教案文档，返回文件路径"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)

    subject = intent.get("subject", "课程")
    topic = intent.get("topic", "教学课题")
    grade = intent.get("grade", "")
    duration = intent.get("duration", "45分钟")
    objectives = intent.get("objectives", [])
    knowledge_points = intent.get("knowledge_points", [])
    key_points = intent.get("key_points", [])
    difficult_points = intent.get("difficult_points", [])
    methods = intent.get("teaching_methods", ["讲授法", "讨论法"])
    activities = intent.get("activities", [])
    student_profile = intent.get("student_profile")

    # ==== 标题 ====
    title = doc.add_heading(f'《{topic}》教学设计', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息表
    doc.add_heading('一、基本信息', level=1)
    info_table = doc.add_table(rows=4, cols=4, style='Table Grid')
    info_data = [
        ["学科", subject, "年级", grade],
        ["课题", topic, "课时", str(duration)],
        ["教师", "________", "日期", "________"],
        ["教材版本", "________", "教学资源", "PPT课件、多媒体素材"],
    ]
    for i, row_data in enumerate(info_data):
        for j, cell_text in enumerate(row_data):
            info_table.rows[i].cells[j].text = cell_text

    doc.add_paragraph()  # 空行

    # ==== 教学目标 ====
    doc.add_heading('二、教学目标', level=1)
    if objectives:
        for obj in objectives:
            doc.add_paragraph(f"• {obj}", style='List Bullet')
    else:
        doc.add_paragraph("• 知识与技能：掌握核心概念和基本原理")
        doc.add_paragraph("• 过程与方法：通过探究学习培养分析能力")
        doc.add_paragraph("• 情感态度与价值观：培养科学精神和创新意识")

    # ==== 教学重点难点 ====
    doc.add_heading('三、教学重点与难点', level=1)
    doc.add_heading('教学重点：', level=2)
    if key_points:
        for p in key_points:
            doc.add_paragraph(f"• {p}", style='List Bullet')
    else:
        doc.add_paragraph("• 核心概念的理解与掌握", style='List Bullet')

    doc.add_heading('教学难点：', level=2)
    if difficult_points:
        for p in difficult_points:
            doc.add_paragraph(f"• {p}", style='List Bullet')
    else:
        doc.add_paragraph("• 抽象概念的形象化理解", style='List Bullet')

    # ==== 教学方法 ====
    doc.add_heading('四、教学方法', level=1)
    method_text = "、".join(methods) if methods else "讲授法、讨论法、案例分析法"
    doc.add_paragraph(f"本课采用{method_text}等多种教学方法，注重启发式教学和学生参与。")

    # ==== 教学过程 ====
    doc.add_heading('五、教学过程', level=1)

    # 生成详细教学过程
    process_prompt = f"""请为《{topic}》（学科：{subject}，课时：{duration}）撰写详细的教学过程。
知识点：{', '.join(knowledge_points) if knowledge_points else '待定'}
重点：{', '.join(key_points) if key_points else '待定'}
难点：{', '.join(difficult_points) if difficult_points else '待定'}

用中文撰写，内容要充实、可操作。"""

    # 生成详细教学过程（优先LLM，失败回退模板）
    try:
        process_text = generate_content(process_prompt, context, student_profile)
        if not process_text or len(process_text) < 50:
            raise ValueError("LLM返回内容过短")
    except Exception:
        process_text = _default_teaching_process(topic, duration)

    # 分段处理
    sections = process_text.split("\n\n")
    for section in sections:
        section = section.strip()
        if not section:
            continue
        lines = section.split("\n")
        if lines:
            first_line = lines[0].strip()
            # 检测是否为标题行
            if any(first_line.startswith(prefix) for prefix in
                   ["1.", "2.", "3.", "4.", "一", "二", "三", "四", "#", "导入", "新课", "巩固", "课堂"]):
                # 清理数字标记
                clean_title = first_line.lstrip("0123456789. #（）()").strip()
                if clean_title:
                    doc.add_heading(clean_title[:50], level=2)
                # 添加剩余内容
                for line in lines[1:]:
                    if line.strip():
                        doc.add_paragraph(line.strip())
            else:
                doc.add_paragraph(section)

    # ==== 课堂活动设计 ====
    doc.add_heading('六、课堂活动设计', level=1)
    if activities:
        for i, act in enumerate(activities):
            doc.add_heading(f'活动{i+1}：{act}', level=2)
            try:
                act_prompt = f"""请为《{topic}》课堂活动【{act}】撰写详细的活动方案（学科：{subject}，年级：{grade}）。
包含：活动目标、活动形式、具体步骤、预计时间。语言简洁可操作。"""
                act_content = generate_content(act_prompt, "", student_profile)
                doc.add_paragraph(act_content)
            except Exception:
                doc.add_paragraph(f"活动目标：配合{act}相关知识点，提升学生参与度。")
                doc.add_paragraph("活动形式：小组合作 / 全班互动")
                doc.add_paragraph("预计时间：5-8分钟")
    else:
        doc.add_paragraph("• 小组讨论：围绕核心知识点展开小组讨论")
        doc.add_paragraph("• 随堂提问：通过提问检查学生理解程度")
        doc.add_paragraph("• 案例分析：结合生活实际进行案例分析")

    # ==== 课后作业 ====
    doc.add_heading('七、课后作业', level=1)
    doc.add_paragraph("1. 复习本节课知识点，完成课后练习题")
    doc.add_paragraph(f"2. 思考题：结合{topic}的相关知识，分析一个实际案例")
    doc.add_paragraph("3. 预习下一节课内容，记录疑问")

    # ==== 教学反思（留白）====
    doc.add_heading('八、教学反思', level=1)
    doc.add_paragraph("（课后填写）")
    doc.add_paragraph("本节课教学效果：________________________________")
    doc.add_paragraph("需要改进之处：________________________________")
    doc.add_paragraph("学生反馈情况：________________________________")

    # 保存文件
    filename = f"教案_{topic}_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)

    return filename


def _default_teaching_process(topic: str, duration: str) -> str:
    """默认教学过程模板"""
    return f"""1. 导入新课（约5分钟）
通过生活中的实际案例引入《{topic}》的话题，激发学生兴趣。
提出引导性问题，让学生带着问题进入新课学习。

2. 新课讲授（约25分钟）
按照知识点的逻辑顺序，由浅入深进行讲解。
每个知识点配合具体案例或数据，帮助学生理解。
穿插课堂提问，及时了解学生掌握情况。
对重点内容进行强调，对难点内容进行多角度解读。

3. 巩固练习（约10分钟）
布置2-3道随堂练习题，覆盖本节课核心知识点。
学生独立完成或小组讨论后，请学生上台展示解题思路。
教师点评，纠正共性错误。

4. 课堂小结（约5分钟）
回顾本节课的知识框架，强调重点难点。
布置课后作业，预告下节课内容。"""


def _generate_quiz_data(topic: str, subject: str, knowledge_points: list, student_profile: dict = None) -> str:
    """生成互动问答数据（优先LLM，失败回退模板），返回 JS 数组的 JSON 字符串"""
    try:
        kp_text = "、".join(knowledge_points[:3]) if knowledge_points else topic
        prompt = f"""请为《{topic}》（学科：{subject}，涉及知识点：{kp_text}）生成3道选择题，用于课堂互动问答。

要求：
- 每道题4个选项（A/B/C/D），只有一个正确答案
- 题目难度递进：第1题基础概念，第2题理解应用，第3题综合分析
- 选项要有迷惑性，错误选项不能太明显
- 每道题的题干和选项都要紧扣知识点

返回严格的JSON数组格式（不要包含markdown代码块标记）：
[
  {{"question": "题目1", "options": ["A. 选项", "B. 选项", "C. 选项", "D. 选项"], "correct": 0}},
  {{"question": "题目2", "options": ["A. 选项", "B. 选项", "C. 选项", "D. 选项"], "correct": 1}},
  {{"question": "题目3", "options": ["A. 选项", "B. 选项", "C. 选项", "D. 选项"], "correct": 2}}
]"""

        llm_output = generate_content(prompt, "", student_profile)
        # 清理可能的 markdown 标记
        llm_output = llm_output.strip()
        if llm_output.startswith("```json"):
            llm_output = llm_output[7:]
        if llm_output.startswith("```"):
            llm_output = llm_output[3:]
        if llm_output.endswith("```"):
            llm_output = llm_output[:-3]
        llm_output = llm_output.strip()
        
        quiz_list = json.loads(llm_output)
        if not isinstance(quiz_list, list) or len(quiz_list) < 2:
            raise ValueError("LLM返回题目不足")
        # 验证每道题结构
        for q in quiz_list:
            if not all(k in q for k in ("question", "options", "correct")):
                raise ValueError("题目格式不正确")
        return json.dumps(quiz_list, ensure_ascii=False)
    except Exception:
        # 回退模板
        quiz_list = [
            {
                "question": f"{topic}的核心概念是什么？",
                "options": ["A. 知识的简单记忆", "B. 理解原理并能应用", "C. 只记住公式即可", "D. 不需要理解"],
                "correct": 1
            },
            {
                "question": f"学习{topic}最重要的方法是？",
                "options": ["A. 死记硬背", "B. 理解+实践", "C. 只看不练", "D. 抄袭作业"],
                "correct": 1
            },
            {
                "question": f"以下哪个是{topic}的实际应用？",
                "options": ["A. 与实际无关", "B. 解决现实问题", "C. 仅用于考试", "D. 无实际用途"],
                "correct": 1
            }
        ]
        return json.dumps(quiz_list, ensure_ascii=False)


def create_html_interactive(intent: dict) -> str:
    """生成HTML互动小游戏/动画页面"""
    topic = intent.get("topic", "知识")
    subject = intent.get("subject", "课程")
    student_profile = intent.get("student_profile")
    knowledge_points = intent.get("knowledge_points", [])

    # 生成互动问答数据（优先LLM，失败回退模板）
    quiz_data_json = _generate_quiz_data(topic, subject, knowledge_points, student_profile)

    html_content = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{topic} - Interactive Quiz</title>
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{
    font-family: -apple-system, 'SF Pro Display', 'Helvetica Neue', 'PingFang SC', sans-serif;
    background: #FBFBFA;
    min-height: 100vh;
    display: flex;
    justify-content: center;
    align-items: center;
    padding: 24px;
    color: #1A1A1A;
    -webkit-font-smoothing: antialiased;
}}
.container {{
    background: #FFFFFF;
    border-radius: 12px;
    padding: 40px 44px;
    max-width: 640px;
    width: 100%;
    border: 1px solid rgba(0,0,0,0.06);
}}
h1 {{
    text-align: center;
    color: #1A1A1A;
    margin-bottom: 4px;
    font-size: 24px;
    font-weight: 600;
    letter-spacing: -0.02em;
}}
.subject-tag {{
    text-align: center;
    color: #6E6E6E;
    margin-bottom: 32px;
    font-size: 13px;
    letter-spacing: 0.01em;
}}
.question-box {{
    background: #FBFBFA;
    border-radius: 8px;
    padding: 24px;
    margin-bottom: 20px;
    border: 1px solid rgba(0,0,0,0.06);
}}
.question-box h2 {{
    color: #1A1A1A;
    margin-bottom: 20px;
    font-size: 18px;
    font-weight: 500;
    letter-spacing: -0.01em;
}}
.options {{ display: flex; flex-direction: column; gap: 8px; }}
.option {{
    padding: 13px 16px;
    background: #FFFFFF;
    border: 1px solid rgba(0,0,0,0.08);
    border-radius: 6px;
    cursor: pointer;
    transition: all 0.15s ease;
    font-size: 14px;
    color: #1A1A1A;
    letter-spacing: -0.01em;
}}
.option:hover {{ border-color: rgba(0,0,0,0.2); background: #F6F6F4; }}
.option.correct {{ border-color: #2E6B35; background: #EDF5EE; color: #2E6B35; }}
.option.wrong {{ border-color: #8B3030; background: #FEF0F0; color: #8B3030; }}
.feedback {{
    margin-top: 14px;
    padding: 12px 14px;
    border-radius: 6px;
    display: none;
    font-size: 13px;
    letter-spacing: -0.01em;
}}
.feedback.show {{ display: block; }}
.feedback.correct {{ background: #EDF5EE; color: #2E6B35; border: 1px solid rgba(46,107,53,0.12); }}
.feedback.wrong {{ background: #FEF0F0; color: #8B3030; border: 1px solid rgba(139,48,48,0.12); }}
.next-btn {{
    display: block;
    width: 100%;
    padding: 12px;
    background: #1A1A1A;
    color: #FFFFFF;
    border: none;
    border-radius: 6px;
    font-size: 14px;
    font-weight: 500;
    cursor: pointer;
    margin-top: 20px;
    transition: background 0.15s ease;
    font-family: inherit;
    letter-spacing: -0.01em;
}}
.next-btn:hover {{ background: #333333; }}
.next-btn:active {{ transform: scale(0.98); }}
.score {{
    text-align: center;
    font-size: 13px;
    color: #6E6E6E;
    margin-top: 14px;
}}
</style>
</head>
<body>
<div class="container">
    <h1>{topic}</h1>
    <p class="subject-tag">{subject} · 互动问答</p>
    <div class="question-box">
        <h2 id="question">点击"开始答题"按钮</h2>
        <div class="options" id="options"></div>
        <div class="feedback" id="feedback"></div>
    </div>
    <button class="next-btn" id="nextBtn" onclick="startQuiz()">开始答题</button>
    <div class="score" id="score"></div>
</div>

<script>
const quizData = {quiz_data_json};

let currentQ = 0, score = 0, answered = false;

function startQuiz() {{
    currentQ = 0; score = 0; answered = false;
    document.getElementById('score').textContent = '';
    document.getElementById('nextBtn').textContent = '下一题';
    document.getElementById('nextBtn').onclick = nextQuestion;
    loadQuestion();
}}

function loadQuestion() {{
    answered = false;
    const q = quizData[currentQ];
    document.getElementById('question').textContent = `Q${{currentQ+1}}. ${{q.question}}`;
    const fb = document.getElementById('feedback');
    fb.classList.remove('show', 'correct', 'wrong');
    const optionsDiv = document.getElementById('options');
    optionsDiv.innerHTML = q.options.map((opt, i) =>
        `<div class="option" onclick="checkAnswer(${{i}})">${{opt}}</div>`
    ).join('');
    if (currentQ >= quizData.length - 1) {{
        document.getElementById('nextBtn').textContent = '查看结果';
    }}
}}

function checkAnswer(index) {{
    if (answered) return;
    answered = true;
    const q = quizData[currentQ];
    const options = document.querySelectorAll('.option');
    const fb = document.getElementById('feedback');
    if (index === q.correct) {{
        options[index].classList.add('correct');
        fb.textContent = '回答正确！';
        fb.classList.add('show', 'correct');
        score++;
    }} else {{
        options[index].classList.add('wrong');
        options[q.correct].classList.add('correct');
        fb.textContent = '再想想，正确答案是 ' + q.options[q.correct].charAt(0);
        fb.classList.add('show', 'wrong');
    }}
}}

function nextQuestion() {{
    if (!answered && currentQ > 0) return;
    currentQ++;
    if (currentQ >= quizData.length) {{
        document.getElementById('question').textContent = '答题完成！';
        document.getElementById('options').innerHTML = '';
        document.getElementById('feedback').classList.remove('show', 'correct', 'wrong');
        document.getElementById('score').textContent = `得分: ${{score}} / ${{quizData.length}}`;
        document.getElementById('nextBtn').textContent = '重新开始';
        document.getElementById('nextBtn').onclick = startQuiz;
        currentQ = quizData.length;
    }} else {{
        loadQuestion();
        document.getElementById('score').textContent = `进度: ${{currentQ+1}}/${{quizData.length}} | 得分: ${{score}}`;
    }}
}}
</script>
</body>
</html>'''

    filename = f"互动问答_{topic}_{uuid.uuid4().hex[:8]}.html"
    filepath = os.path.join(OUTPUT_DIR, filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html_content)

    return filename
